from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from core_logic import (  # noqa: E402
    DocumentExtractionError,
    build_proposed_filename,
    extract_document,
    get_basic_document_details,
    get_document_details,
    normalize_document_details,
    validate_document_filename,
    validate_filename_template,
)
from document_types import is_processable_document  # noqa: E402


class CoreLogicTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.root = Path(self.tempdir.name)
        self.docx_path = self.root / "report.docx"
        document = Document()
        document.core_properties.title = "A Practical Technical Report"
        document.core_properties.author = "Jane Example"
        document.add_heading("A Practical Technical Report", level=1)
        document.add_paragraph("Technical Report No. 42")
        document.add_paragraph("Published 2024 by Example Institute.")
        document.add_table(rows=1, cols=2).rows[0].cells[0].text = "Summary"
        document.tables[0].rows[0].cells[1].text = "Local extraction works."
        document.save(self.docx_path)

    def tearDown(self):
        self.tempdir.cleanup()

    def test_docx_is_extracted_locally(self):
        extraction = extract_document(self.docx_path)

        self.assertEqual(extraction.document_format, "DOCX")
        self.assertIn("Technical Report", extraction.text)
        self.assertEqual(extraction.metadata["title"], "A Practical Technical Report")

    def test_basic_docx_details_are_extension_safe_without_a_fake_confidence_score(self):
        details = get_basic_document_details(self.docx_path)

        self.assertEqual(details["source"], "Basic")
        self.assertEqual(details["document_type"], "technical_report")
        self.assertEqual(details["evidence_label"], "Suggested from local document information")
        self.assertNotIn("confidence", details)
        proposal = build_proposed_filename(details, self.docx_path.suffix)
        self.assertTrue(proposal.endswith(".docx"))
        self.assertNotIn(".pdf", proposal)

    def test_docx_ai_is_local_by_default(self):
        details = get_document_details(self.docx_path, "not-a-real-key", allow_cloud_ai=False)

        self.assertEqual(details["source"], "Basic")

    def test_filename_validation_preserves_original_type(self):
        self.assertEqual(validate_document_filename("Edited Report", ".docx"), "Edited Report.docx")
        self.assertEqual(validate_document_filename("Edited.Report", ".docx"), "Edited.Report.docx")
        self.assertEqual(validate_document_filename("Slide Deck", ".pptx"), "Slide Deck.pptx")
        self.assertEqual(validate_document_filename("Legacy Slides", ".ppt"), "Legacy Slides.ppt")
        with self.assertRaisesRegex(ValueError, "original .docx extension"):
            validate_document_filename("Edited Report.pdf", ".docx")
        with self.assertRaisesRegex(ValueError, "original .pptx extension"):
            validate_document_filename("Slide Deck.pdf", ".pptx")

    def test_normalizer_does_not_treat_string_false_as_true(self):
        details = normalize_document_details(
            {
                "title": "Example",
                "primary_creator": "Jane Example",
                "year": "2025",
                "document_type": "report",
                "is_multiple_creators": "false",
                "confidence": "0.91",
                "needs_review": "false",
            },
            self.docx_path,
            source="AI",
        )

        self.assertFalse(details["is_multiple_creators"])
        self.assertFalse(details["needs_review"])
        self.assertEqual(details["document_type"], "technical_report")
        self.assertNotIn("confidence", details)

    def test_office_lock_file_is_never_accepted(self):
        lock_path = self.root / "~$report.docx"
        lock_path.write_bytes(b"not a document")

        self.assertFalse(is_processable_document(lock_path))
        with self.assertRaises(DocumentExtractionError):
            extract_document(lock_path)

    def test_citation_styles_use_verified_journal_fields_and_keep_extension(self):
        details = {
            "title": "A Practical Example Study",
            "primary_creator": "Jane Doe",
            "author": "Doe",
            "year": "2024",
            "document_type": "journal_article",
            "venue_or_publisher": "Journal of Example Research",
            "journal": "Journal of Example Research",
            "journal_abbreviation": "J Example Res",
            "volume": "12",
            "issue": "3",
            "is_multiple_creators": True,
        }

        self.assertEqual(
            build_proposed_filename(details, ".pdf", filename_format="journal_compact"),
            "Doe_et_al_J_Example_Res_2024.pdf",
        )
        self.assertEqual(
            build_proposed_filename(details, ".docx", filename_format="journal_detailed"),
            "Doe_et_al_Journal_of_Example_Research_12_3_2024.docx",
        )
        self.assertEqual(
            build_proposed_filename(details, ".pdf", filename_format="author_year_title"),
            "Doe_et_al_2024_A_Practical_Example_Study.pdf",
        )

    def test_journal_abbreviation_is_never_invented_and_custom_templates_are_safe(self):
        details = {
            "title": "A Practical Example Study",
            "primary_creator": "Jane Doe",
            "author": "Doe",
            "year": "2024",
            "document_type": "journal_article",
            "venue_or_publisher": "Journal of Example Research",
            "journal": "Journal of Example Research",
            "is_multiple_creators": True,
        }

        self.assertEqual(
            build_proposed_filename(details, ".pdf", filename_format="journal_compact"),
            "Jane_Doe_et_al_Journal_of_Example_Research_2024.pdf",
        )
        self.assertEqual(
            build_proposed_filename(
                details,
                ".docx",
                filename_format="custom",
                custom_template="{author_last}_{volume}_{issue}_{year}_{title}",
            ),
            "Doe_2024_A_Practical_Example_Study.docx",
        )
        with self.assertRaisesRegex(ValueError, "Unsupported template token"):
            validate_filename_template("{run_command}")
        with self.assertRaisesRegex(ValueError, "original extension"):
            validate_filename_template("{title}.pdf")

    def test_normalizer_retains_bibliographic_fields_for_filename_styles(self):
        details = normalize_document_details(
            {
                "title": "Example",
                "primary_creator": "Jane Doe",
                "year": "2025",
                "document_type": "journal_article",
                "journal": "Journal of Examples",
                "journal_abbreviation": "J Examples",
                "volume": "19",
                "issue": "2",
                "confidence": 0.95,
            },
            self.docx_path,
            source="AI",
        )

        self.assertEqual(details["journal_abbreviation"], "J Examples")
        self.assertEqual(details["volume"], "19")
        self.assertEqual(details["issue"], "2")


if __name__ == "__main__":
    unittest.main()
